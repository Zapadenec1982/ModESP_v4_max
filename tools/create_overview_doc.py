#!/usr/bin/env python3
"""
Генерує презентаційний документ ModESP v4 у форматі .docx
Запуск: python tools/create_overview_doc.py
Результат: docs/ModESP_v4_overview_ua.docx
"""

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
import os

def set_cell_shading(cell, color_hex):
    """Встановити фон комірки таблиці."""
    from docx.oxml.ns import qn
    from lxml import etree
    shading = etree.SubElement(cell._tc.get_or_add_tcPr(), qn('w:shd'))
    shading.set(qn('w:fill'), color_hex)
    shading.set(qn('w:val'), 'clear')

def add_bullet(doc, text, bold_prefix=None, level=0):
    """Додати bullet point з опціональним жирним префіксом."""
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(1)
    if level > 0:
        p.paragraph_format.left_indent = Cm(1.5 * level)
    if bold_prefix:
        run_bold = p.add_run(bold_prefix)
        run_bold.bold = True
        run_bold.font.size = Pt(10.5)
        run_normal = p.add_run(text)
        run_normal.font.size = Pt(10.5)
    else:
        run = p.add_run(text)
        run.font.size = Pt(10.5)
    return p

def add_body(doc, text):
    """Додати абзац основного тексту."""
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    for run in p.runs:
        run.font.size = Pt(10.5)
    return p

def add_body_rich(doc, parts):
    """Додати абзац з mix жирного та звичайного тексту.
    parts = [("text", bold), ("text2", bold2), ...]
    """
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    for text, bold in parts:
        run = p.add_run(text)
        run.bold = bold
        run.font.size = Pt(10.5)
    return p

def create_comparison_table(doc):
    """Таблиця порівняння ModESP vs Danfoss/Dixell vs Arduino."""
    table = doc.add_table(rows=11, cols=4)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ['Можливість', 'ModESP v4', 'Danfoss / Dixell', 'DIY Arduino']
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9.5)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(cell, '2B5797')

    rows_data = [
        ['Вартість контролера',      '~$5-8 (ESP32)',   '$30-80+',          '$5-10'],
        ['Холодильна логіка',        'Повна',           'Повна',            'Відсутня'],
        ['Веб-інтерфейс',            'Так (WiFi)',      'Ні (тільки LCD)',  'Ні'],
        ['Віддалений моніторинг',    'MQTT + WebUI',    'Пропрієтарний',    'Частково'],
        ['OTA оновлення',            'Так + rollback',  'Ні (заміна плати)','Ні (USB)'],
        ['Захист компресора',        '10 моніторів',    '6-8 моніторів',    'Відсутній'],
        ['Логування даних',          '6 каналів',       'Опціонально',      'Ні'],
        ['Нічний режим',             '4 режими',        '2 режими',         'Ні'],
        ['Кастомізація',             'Повна (JSON)',    'Обмежена',         'Потребує коду'],
        ['Відкритий код',            'Так',             'Ні',               'Залежить'],
    ]

    for row_idx, row_data in enumerate(rows_data):
        for col_idx, val in enumerate(row_data):
            cell = table.rows[row_idx + 1].cells[col_idx]
            cell.text = val
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
                if col_idx == 0:
                    for run in p.runs:
                        run.bold = True
            if row_idx % 2 == 0:
                set_cell_shading(cell, 'E8F0FE')

    return table

def create_modules_table(doc):
    """Таблиця модулів системи."""
    table = doc.add_table(rows=6, cols=3)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ['Модуль', 'Призначення', 'Ключові можливості']
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9.5)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(cell, '2B5797')

    rows_data = [
        ['Equipment\nManager',
         'Управління обладнанням\n(реле, датчики)',
         'Арбітраж пріоритетів, інтерлоки безпеки,\nEMA-фільтр температур, 6 типів драйверів'],
        ['Thermostat',
         'Регулювання\nтемператури',
         'Асиметричний диференціал, 3 режими вентилятора,\nнічний режим (4 варіанти), safety run при відмові датчика'],
        ['Defrost',
         'Відтайка\nвипарника',
         '7-фазний цикл, 3 типи (природна/тен/гарячий газ),\n4 ініціації, оптимізація (skip якщо чистий)'],
        ['Protection',
         'Захист обладнання\nта аварії',
         '10 моніторів аварій, захист компресора,\nмоторгодини, діагностика, auto/manual reset'],
        ['DataLogger',
         'Логування даних\nта аналітика',
         '6 каналів температур, 10 типів подій,\nграфік у WebUI, CSV експорт'],
    ]

    for row_idx, row_data in enumerate(rows_data):
        for col_idx, val in enumerate(row_data):
            cell = table.rows[row_idx + 1].cells[col_idx]
            cell.text = val
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
                if col_idx == 0:
                    for run in p.runs:
                        run.bold = True
            if row_idx % 2 == 0:
                set_cell_shading(cell, 'E8F0FE')

    return table


def main():
    doc = Document()

    # --- Стилі ---
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10.5)
    style.paragraph_format.space_after = Pt(4)

    for level in range(1, 4):
        heading_style = doc.styles[f'Heading {level}']
        heading_style.font.name = 'Calibri'
        heading_style.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)

    # --- Обкладинка ---
    for _ in range(4):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('ModESP v4')
    run.bold = True
    run.font.size = Pt(32)
    run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('Модульна платформа керування\nхолодильним обладнанням')
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x4A, 0x4A, 0x4A)

    doc.add_paragraph()

    tagline = doc.add_paragraph()
    tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = tagline.add_run('Професійна надійність. Сучасні технології. Доступна вартість.')
    run.font.size = Pt(12)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    for _ in range(8):
        doc.add_paragraph()

    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_p.add_run('Березень 2026')
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_page_break()

    # ========================================================
    # 1. ЩО ТАКЕ ModESP
    # ========================================================
    doc.add_heading('Що таке ModESP', level=1)

    add_body(doc,
        'ModESP v4 — це модульний програмний фреймворк для промислових контролерів '
        'холодильного обладнання на базі мікроконтролера ESP32. '
        'Проект створений як доступна та гнучка альтернатива дорогим пропрієтарним '
        'контролерам від Danfoss, Dixell, Carel та інших виробників.'
    )

    add_body(doc,
        'Основна ідея: один і той самий firmware може обслуговувати різні типи обладнання — '
        'від простих морозильних скринь до складних мультизонних систем. '
        'Конфігурація під конкретний тип обладнання відбувається через JSON-файли, '
        'без необхідності змінювати програмний код.'
    )

    add_body_rich(doc, [
        ('Апаратна база: ', True),
        ('ESP32-WROOM-32 — широко доступний мікроконтролер з WiFi та Bluetooth, '
         'вартість якого складає $3-5 в роздріб (при серійних закупівлях — значно менше). '
         'Для порівняння: один контролер Danfoss ERC 213 коштує $30-80 і більше.', False),
    ])

    # ========================================================
    # 2. ПРОБЛЕМА РИНКУ
    # ========================================================
    doc.add_heading('Чому це актуально', level=1)

    add_body(doc,
        'Сучасні виробники холодильного обладнання стикаються з кількома викликами, '
        'пов\'язаними з електронними контролерами:'
    )

    add_bullet(doc, ' — стандартні контролери Danfoss, Dixell, Carel '
        'додають $30-80+ до собівартості кожної одиниці обладнання. '
        'При серійному виробництві десятків тисяч одиниць на рік це формує '
        'значну статтю витрат.',
        'Вартість.')

    add_bullet(doc, ' — пропрієтарні контролери мають фіксований набір функцій. '
        'Додати нову можливість (наприклад, IoT-моніторинг чи специфічну логіку відтайки) '
        'неможливо без заміни контролера або дорогої кастомної розробки від виробника контролера.',
        'Закритість.')

    add_bullet(doc, ' — в умовах війни ланцюги поставок електроніки з Європи '
        'стали менш стабільними. Залежність від одного постачальника контролерів '
        'створює ризики для безперервності виробництва.',
        'Ланцюги поставок.')

    add_bullet(doc, ' — європейський та глобальний ринок все більше вимагає '
        '"розумне" обладнання з можливістю віддаленого моніторингу, '
        'аналітики споживання та предиктивного обслуговування. '
        'Базові контролери цього не забезпечують.',
        'IoT-тренд.')

    # ========================================================
    # 3. ЩО ВМІЄ ModESP ЗАРАЗ
    # ========================================================
    doc.add_heading('Поточний стан проекту', level=1)

    add_body_rich(doc, [
        ('ModESP v4 пройшов ', False),
        ('14 сесій активної розробки ', True),
        ('і на сьогодні є повністю функціональною системою, '
         'яка пройшла тестування на реальному обладнанні (плата KC868-A6, 30+ годин безперервної роботи).', False),
    ])

    doc.add_heading('Бізнес-модулі', level=2)

    add_body(doc, 'Система складається з п\'яти основних модулів, кожен з яких відповідає '
        'за свою частину холодильного циклу:')

    create_modules_table(doc)

    doc.add_paragraph()  # spacing

    doc.add_heading('Веб-інтерфейс', level=2)

    add_body(doc,
        'Кожен контролер має вбудований веб-сервер, доступний через WiFi. '
        'Для налаштування та моніторингу достатньо смартфона або ноутбука — '
        'жодного додаткового ПЗ не потрібно.'
    )

    add_bullet(doc, 'Dashboard з реальними показниками температур, станом реле та аварій')
    add_bullet(doc, 'Сторінки налаштувань для кожного модуля (термостат, відтайка, захист, логування)')
    add_bullet(doc, 'Графік температур з історією (24/48 годин) та подіями')
    add_bullet(doc, 'Конфігурація обладнання (прив\'язка датчиків та реле)')
    add_bullet(doc, 'Світла та темна теми, українська та англійська мови')
    add_bullet(doc, 'Адаптивний дизайн для мобільних пристроїв')

    doc.add_heading('Комунікації та інтеграція', level=2)

    add_bullet(doc, ' — підключення через WiFi, передача даних за протоколом MQTT '
        '(стандарт для IoT). Можливість інтеграції з будь-якою системою моніторингу.',
        'MQTT.')
    add_bullet(doc, ' — оновлення прошивки контролерів через WiFi, без фізичного доступу. '
        'Автоматичний відкат до попередньої версії у разі проблеми.',
        'OTA оновлення.')
    add_bullet(doc, ' — повний REST API (21 endpoint) для інтеграції з зовнішніми системами, '
        'WebSocket для real-time даних.',
        'HTTP API.')

    doc.add_heading('Порівняння з існуючими рішеннями', level=2)

    create_comparison_table(doc)

    doc.add_paragraph()  # spacing

    # ========================================================
    # 4. АРХІТЕКТУРНІ ПЕРЕВАГИ
    # ========================================================
    doc.add_heading('Інженерний підхід', level=1)

    add_body(doc,
        'ModESP побудований на принципах промислової надійності, а не "домашнього DIY". '
        'Ось ключові архітектурні рішення:'
    )

    add_bullet(doc, ' — Equipment Manager арбітрує всі запити до обладнання: '
        'захист має найвищий пріоритет, потім відтайка, потім термостат. '
        'Компресор та нагрівач ніколи не можуть бути увімкнені одночасно (інтерлок безпеки).',
        'Арбітраж безпеки.')

    add_bullet(doc, ' — описані в JSON-маніфестах, модулі можна додавати '
        'та конфігурувати без зміни основного коду. Один firmware — різні конфігурації.',
        'Модульність.')

    add_bullet(doc, ' — C++17 з ETL (Embedded Template Library), '
        'zero heap allocation в основному циклі — програма не може "зависнути" через '
        'фрагментацію пам\'яті, як це буває в Arduino-проектах.',
        'Промислова стабільність.')

    add_bullet(doc, ' — 354 автоматичних тести перевіряють логіку кожного модуля. '
        'Кожна зміна в коді проходить автоматичну верифікацію.',
        'Тестування.')

    add_bullet(doc, ' — anti-short-cycle protection (мінімальний час ON/OFF), '
        'відстеження кількості пусків за годину, duty cycle, моторгодини — '
        'все що потрібно для подовження ресурсу компресора.',
        'Захист компресора.')

    # ========================================================
    # 5. ПЕРСПЕКТИВИ ДЛЯ ВИРОБНИКА
    # ========================================================
    doc.add_heading('Перспективи для виробника обладнання', level=1)

    doc.add_heading('Зниження собівартості', level=2)
    add_body(doc,
        'ESP32 коштує $3-5 (роздріб) або $1.5-3 при оптових закупівлях. '
        'Порівняно з $30-80 за контролер Danfoss/Dixell, економія на кожній одиниці '
        'обладнання може складати $25-75. При річному виробництві 50 000 одиниць '
        'це від $1.25 до $3.75 млн економії тільки на контролерах.'
    )

    doc.add_heading('Повний контроль над продуктом', level=2)
    add_body(doc,
        'Виробник отримує повний контроль над логікою свого обладнання: '
        'може додавати унікальні алгоритми, специфічні для своїх продуктових ліній, '
        'оптимізувати під конкретні компресори та хладагенти, '
        'створювати фірмовий веб-інтерфейс з власним брендингом.'
    )

    doc.add_heading('IoT як конкурентна перевага', level=2)
    add_body(doc,
        'Вбудований WiFi та MQTT з коробки дозволяють виробнику запропонувати '
        'клієнтам (торговим мережам, ресторанам, складам) систему віддаленого '
        'моніторингу всього парку обладнання. Це особливо актуально для виходу '
        'на європейський ринок, де IoT-функціональність стає стандартною вимогою.'
    )

    doc.add_heading('Незалежність від постачальників', level=2)
    add_body(doc,
        'ESP32 виробляється компанією Espressif (Китай/Індія) і доступний через '
        'десятки дистриб\'юторів по всьому світу. На відміну від спеціалізованих '
        'контролерів Danfoss/Dixell, замовити ESP32 можна у будь-який момент, '
        'від будь-якого постачальника, з мінімальним lead time.'
    )

    doc.add_heading('Масштабованість рішення', level=2)
    add_body(doc,
        'Один і той самий firmware підтримує різні апаратні конфігурації: '
        'від простої плати з GPIO-реле до I2C-розширювачів (KC868-A6 з 6 реле та 6 входами). '
        'Це означає що одна розробка покриває всю лінійку продуктів — '
        'від базових моделей до преміум-сегменту.'
    )

    # ========================================================
    # 6. ПЕРСПЕКТИВИ ДЛЯ СЕРВІСУ
    # ========================================================
    doc.add_heading('Перспективи для сервісної служби', level=1)

    add_body(doc,
        'Для сервісного підрозділу ModESP відкриває принципово нові можливості:'
    )

    add_bullet(doc, ' — доступ до стану обладнання через браузер або MQTT-dashboard '
        'з будь-якого місця. Перевірити температури, стан реле, активні аварії, '
        'час роботи компресора — без виїзду на об\'єкт.',
        'Віддалена діагностика.')

    add_bullet(doc, ' — 6-канальний логер записує температури повітря, випарника, '
        'конденсатора, а також події (пуск/зупинка компресора, початок/кінець відтайки, '
        'аварії, відчинення дверей). Ці дані доступні як графік у WebUI '
        'та як CSV-файл для аналізу.',
        'Історія та аналітика.')

    add_bullet(doc, ' — оновлення прошивки через WiFi дозволяє виправляти '
        'проблеми та додавати функціональність без фізичного доступу до обладнання. '
        'Автоматичний rollback гарантує безпеку.',
        'OTA оновлення.')

    add_bullet(doc, ' — контролер відстежує кумулятивні моторгодини, '
        'кількість пусків за годину, duty cycle. Це дозволяє планувати '
        'технічне обслуговування на основі реальних даних, а не календарних інтервалів.',
        'Предиктивне обслуговування.')

    add_bullet(doc, ' — 10 типів аварій з конфігурованими затримками, '
        'автоматичне та ручне скидання. Чіткі діагностичні коди '
        'замість загального "error" на дисплеї.',
        'Прозорі аварії.')

    # ========================================================
    # 7. ДОРОЖНЯ КАРТА
    # ========================================================
    doc.add_heading('Дорожня карта розвитку', level=1)

    add_body(doc, 'Проект активно розвивається. Основні напрямки подальшої роботи:')

    doc.add_heading('Найближча перспектива', level=2)
    add_bullet(doc, 'HTTPS — шифроване з\'єднання з контролером')
    add_bullet(doc, 'LCD/OLED дисплей — для роботи без WiFi (SSD1306)')
    add_bullet(doc, 'Secure Boot — захист прошивки від несанкціонованої зміни')
    add_bullet(doc, 'Розширене тестування на реальному обладнанні')

    doc.add_heading('Середньострокова перспектива', level=2)
    add_bullet(doc, 'Modbus RTU/TCP — інтеграція з існуючими промисловими системами')
    add_bullet(doc, 'Home Assistant — інтеграція з популярною платформою автоматизації')
    add_bullet(doc, 'Fleet Management — backend-dashboard для управління парком обладнання')
    add_bullet(doc, 'Групове OTA оновлення — масове оновлення прошивки')
    add_bullet(doc, 'Сповіщення — email/Telegram при аваріях')

    doc.add_heading('Довгострокова перспектива', level=2)
    add_bullet(doc, 'Мультизонне управління — кілька зон в одному контролері')
    add_bullet(doc, 'PID-регулювання — плавне управління замість on/off')
    add_bullet(doc, 'Каскадне управління — координація кількох контролерів')
    add_bullet(doc, 'Аналітика та звітність — хмарний збір даних, тренди, KPI')
    add_bullet(doc, 'Сертифікація — для специфічних ринків та стандартів')

    # ========================================================
    # 8. ПІДСУМОК
    # ========================================================
    doc.add_heading('Підсумок', level=1)

    add_body(doc,
        'ModESP v4 — це не концепт і не прототип. Це працююча система з повним '
        'циклом холодильного контролю, яка пройшла 30+ годин реальної експлуатації. '
        'На сьогодні реалізовано:'
    )

    add_bullet(doc, '5 бізнес-модулів з повною холодильною логікою')
    add_bullet(doc, '6 драйверів для різних типів обладнання')
    add_bullet(doc, '10 моніторів аварій з захистом компресора')
    add_bullet(doc, 'Сучасний веб-інтерфейс з графіками та логуванням')
    add_bullet(doc, 'MQTT для віддаленого моніторингу')
    add_bullet(doc, 'OTA оновлення з автоматичним rollback')
    add_bullet(doc, '354 автоматичних тестів')
    add_bullet(doc, 'Підтримка 2 апаратних платформ')

    doc.add_paragraph()

    add_body_rich(doc, [
        ('Проект поєднує ', False),
        ('професійну холодильну логіку рівня Danfoss ', True),
        ('з ', False),
        ('сучасними IoT-технологіями ', True),
        ('при вартості комплектуючих у ', False),
        ('десятки разів нижчій ', True),
        ('за існуючі рішення. Це відкриває реальну можливість для виробника обладнання '
         'отримати повний контроль над електронною складовою свого продукту, '
         'одночасно знизивши собівартість та додавши IoT-функціональність.', False),
    ])

    # ========================================================
    # Контакт
    # ========================================================
    doc.add_paragraph()
    doc.add_paragraph()

    separator = doc.add_paragraph()
    separator.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = separator.add_run('— — —')
    run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

    doc.add_paragraph()

    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = contact.add_run('Готовий обговорити деталі та продемонструвати систему в дії.')
    run.font.size = Pt(11)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    # --- Збереження ---
    output_dir = os.path.join(os.path.dirname(os.path.dirname(__file__)), 'docs')
    os.makedirs(output_dir, exist_ok=True)
    output_path = os.path.join(output_dir, 'ModESP_v4_overview_ua.docx')
    doc.save(output_path)
    print(f'Document saved: {output_path}')
    print(f'Size: {os.path.getsize(output_path) / 1024:.1f} KB')


if __name__ == '__main__':
    main()
